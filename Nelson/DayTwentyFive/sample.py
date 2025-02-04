from faker import Faker
from docx import Document
import os

# 创建一个Faker实例
fake = Faker()

# 离职证明模板
def create_resignation_certificate(name, company, date):
    doc = Document()

    # 添加标题
    doc.add_heading('离职证明', 0)

    # 添加正文内容
    doc.add_paragraph(f"兹证明 {name} 在 {company} 工作，并于 {date} 正式离职。")
    doc.add_paragraph(f"{name} 在本公司的表现一直非常优秀，感谢其为公司做出的贡献。")
    doc.add_paragraph(f"公司全体同仁祝愿 {name} 在未来的工作和生活中一切顺利，事业更上一层楼。")

    # 添加公司信息和签名
    doc.add_paragraph(f"公司：{company}")
    doc.add_paragraph(f"签名：_________________")
    doc.add_paragraph(f"日期：{date}")

    return doc

# 批量生成离职证明
def generate_bulk_resignation_certificates(num, output_folder="/Users/nelsonezeanya/Documents/Python_100_days/Nelson/DayTwentyFive"):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for i in range(num):
        name = fake.name()
        company = fake.company()
        date = fake.date()

        # 生成离职证明文档
        doc = create_resignation_certificate(name, company, date)

        # 保存文档
        filename = os.path.join(output_folder, f"离职证明_{i+1}.docx")
        doc.save(filename)

        print(f"生成 {filename}")

# 生成2份离职证明
generate_bulk_resignation_certificates(2)
